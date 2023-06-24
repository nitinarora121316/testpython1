import docx
import pyttsx3

def read_aloud_from_docx(docx_file):
    # Open the Word document
    doc = docx.Document(docx_file)

    # Initialize male and female text-to-speech engines
    engine_male = pyttsx3.init()
    engine_female = pyttsx3.init()

    # Get the available voices
    voices = engine_male.getProperty('voices')

    # Initialize voice indices
    male_voice_index = 0
    female_voice_index = 0

      # Set slower pace for male voice
    engine_male.setProperty('rate', 150)  # Adjust the value to change the pace

    # Set slower pace for female voice
    engine_female.setProperty('rate', 150)  # Adjust the value to change the pace


    # Iterate over paragraphs in a cyclic manner
    paragraph_count = len(doc.paragraphs)
    for i in range(paragraph_count):
        # Get the current paragraph
        paragraph = doc.paragraphs[i]

        # Strip leading/trailing whitespaces
        text = paragraph.text.strip()

        # Check if the paragraph is non-empty
        if text:
            # Determine the voice based on paragraph index
            if i % 2 == 0:
                # Use a male voice
                engine_male.setProperty('voice', voices[male_voice_index % len(voices)].id)
                male_voice_index += 1
                engine_male.say(text)
                engine_male.runAndWait()
            else:
                # Use a female voice
                engine_female.setProperty('voice', voices[female_voice_index % len(voices)].id)
                female_voice_index += 1
                engine_female.say(text)
                engine_female.runAndWait()

# Specify the path to your Word document
word_file = r'C:\Users\asus\Desktop\abc.docx'

# Call the function to read aloud the content of the Word document
read_aloud_from_docx(word_file)
